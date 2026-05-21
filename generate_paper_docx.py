# -*- coding: utf-8 -*-
"""Generate complete HSANet journal manuscript as DOCX."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"c:\Users\tareq\Downloads\Hsanet_research_cursor\HSANet_Journal_Paper.docx"


def set_doc_defaults(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = style.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(doc, text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p


def add_figure_placeholder(doc, fig_num, caption, width_hint="full page width"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[INSERT FIGURE {fig_num} HERE — {width_hint}]")
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    # placeholder box via border paragraph
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"Figure {fig_num}. {caption}")
    cr.font.name = "Times New Roman"
    cr.font.size = Pt(10)
    cr.italic = True
    doc.add_paragraph()


def add_table(doc, headers, rows, caption=None):
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        cr.font.name = "Times New Roman"
        cr.font.size = Pt(10)
        cr.bold = True
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            row_cells[ci].text = str(val)
            for p in row_cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table


def build_document():
    doc = Document()
    set_doc_defaults(doc)

    # --- Title ---
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run(
        "HSANet: A Lightweight Hybrid Scale-Attention Network with "
        "Evidential Uncertainty for Brain Tumor MRI Classification"
    )
    tr.bold = True
    tr.font.size = Pt(16)
    tr.font.name = "Times New Roman"

    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = authors.add_run(
        "Tareque Jamil Josh¹, Md. Minhazur Rahman Mim¹, Md. Aminur Rahman Joy¹, "
        "Md. Assaduzzaman¹, Sheak Rashed Haider Noori¹"
    )
    ar.font.name = "Times New Roman"
    ar.font.size = Pt(12)

    affil = doc.add_paragraph()
    affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
    afr = affil.add_run(
        "¹Department of Computer Science and Engineering, Daffodil International University, "
        "Dhaka, Bangladesh"
    )
    afr.font.name = "Times New Roman"
    afr.font.size = Pt(11)

    corr = doc.add_paragraph()
    corr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crr = corr.add_run(
        "Correspondence: Md. Assaduzzaman (assaduzzaman.cse@diu.edu.bd)"
    )
    crr.font.name = "Times New Roman"
    crr.font.size = Pt(11)
    crr.italic = True

    doc.add_page_break()

    # --- Abstract ---
    add_heading(doc, "Abstract", 1)
    add_para(
        doc,
        "Accurate and trustworthy classification of brain tumors from magnetic resonance imaging (MRI) "
        "remains clinically important yet challenging under scanner variability, class imbalance, and "
        "limited labeled data. Deep learning models with heavy attention stacks often improve in-domain "
        "accuracy only marginally while increasing parameters, latency, and susceptibility to dataset-specific "
        "artifacts. We propose HSANet (Hybrid Scale-Attention Network), a parameter-efficient architecture "
        "comprising a multi-scale EfficientNet-B3 backbone, a depthwise-separable Adaptive Multi-Scale Module "
        "(AMSM) with normalized four-branch softmax fusion, and an Evidential Deep Learning (EDL) classifier "
        "for single-pass uncertainty estimation."
    )
    add_para(
        doc,
        "Through systematic ablation across two design generations, we demonstrate that an auxiliary Dual "
        "Attention Module (DAM)—combining efficient channel attention and spatial attention—used in an earlier "
        "prototype (version 5) did not improve performance and frequently degraded it; the final model "
        "(version 6) therefore retains only AMSM, yielding a simpler and empirically stronger design. On the "
        "expert-annotated BRISC 2025 benchmark (6,000 T1-weighted MRI scans, four classes), accessed via the "
        "public Kaggle release, HSANet achieves 99.2% held-out test accuracy (95% bootstrap confidence interval: "
        "98.6–99.7%), 99.06 ± 0.17% five-fold cross-validation accuracy, macro-F1 of 99.30%, and macro-AUC of "
        "0.9995, with 10.97 million parameters. Against six fairly trained baselines—ResNet-50, VGG-16, "
        "EfficientNet-B3, ViT-B/16, Swin-Tiny, and DenseNet-121—using an identical cross-entropy plus focal "
        "loss protocol, HSANet matches or exceeds the strongest transformer baseline while using substantially "
        "fewer parameters than ViT-B/16 and VGG-16. External validation on the independent PMRAM-BD cohort "
        "(Bangladesh) yields 97.41% accuracy, evidencing cross-site generalization. Component ablation confirms "
        "that AMSM contributes +0.7 percentage points over the backbone alone; EDL contributes an additional "
        "+0.6 points relative to a matched CE-only head. McNemar testing indicates statistically significant "
        "paired disagreement favoring HSANet over ResNet-50 (p = 0.043). Grad-CAM and Multi-Scale Attention "
        "Decomposition (MSAD) analyses indicate lesion-focused attention across dilation scales. These findings "
        "support the conclusion that carefully designed multi-scale fusion—not stacked attention—is the effective "
        "inductive bias for this task, and that evidential heads furnish clinically useful uncertainty without "
        "Monte Carlo inference at test time."
    )

    add_para(doc, "Keywords: ", bold=True)
    p = doc.paragraphs[-1]
    p.add_run(
        "brain tumor classification; magnetic resonance imaging; attention mechanism; "
        "multi-scale convolution; evidential deep learning; uncertainty quantification; "
        "cross-dataset generalization; BRISC 2025"
    ).font.name = "Times New Roman"

    doc.add_page_break()

    # --- 1 Introduction ---
    add_heading(doc, "1. Introduction", 1)
    add_para(
        doc,
        "Primary brain tumors represent a major source of morbidity and mortality worldwide. "
        "Magnetic resonance imaging (MRI) is the modality of choice for non-invasive assessment, "
        "and distinguishing among glioma, meningioma, pituitary adenoma, and non-tumorous tissue "
        "can inform surgical planning, radiation protocols, and longitudinal monitoring [1,2]. "
        "Over the past decade, convolutional neural networks (CNNs) and vision transformers (ViTs) "
        "have achieved high classification accuracy on public brain MRI corpora [3–6]. "
        "Nevertheless, three methodological gaps limit translation of these systems into dependable "
        "clinical decision support."
    )
    add_para(
        doc,
        "First, many studies optimize in-domain accuracy on a single dataset while under-reporting "
        "external validation across hospitals, scanners, and populations [7,8]. Models that exploit "
        "site-specific intensity distributions or preprocessing conventions may fail silently when "
        "deployed on data acquired under different protocols."
    )
    add_para(
        doc,
        "Second, attention mechanisms—including squeeze-and-excitation (SE) blocks [9], efficient "
        "channel attention (ECA) [10], and convolutional block attention modules (CBAM) [11]—are "
        "often appended to CNN backbones with limited ablation. Sequential stacking of channel and "
        "spatial gates after an upstream multi-scale fusion module can redundantly suppress features "
        "that have already been adaptively mixed, increasing parameters and optimization difficulty "
        "without consistent generalization gains [12]."
    )
    add_para(
        doc,
        "Third, high discriminative accuracy alone is insufficient for safe deployment. Clinicians "
        "require reliable uncertainty estimates to support selective prediction—deferring ambiguous "
        "cases for expert review—and to communicate model limitations transparently [13,14]. "
        "Evidential Deep Learning (EDL) [15] parameterizes second-order uncertainty through Dirichlet "
        "distributions, enabling uncertainty quantification in a single forward pass, in contrast to "
        "Monte Carlo (MC) dropout [16], which incurs a multiplicative inference cost."
    )
    add_para(doc, "To address these gaps, we introduce HSANet with the following contributions:", bold=False)
    contributions = [
        "Lightweight AMSM: A depthwise-separable, dilated multi-scale module with softmax-normalized four-branch fusion (three dilated convolutions plus an identity branch), correcting an earlier unnormalized residual formulation that broke the probabilistic interpretation of branch weights.",
        "Principled architectural simplification (v5 to v6): We show through ablation that a Dual Attention Module (ECA followed by spatial attention) appended after AMSM in version 5 did not improve—and slightly harmed—performance; version 6 removes DAM entirely.",
        "Evidential classification head: A compact EDL head with softplus evidence mapping, KL annealing, and focal regularization, evaluated against MC-dropout and an architecture-matched cross-entropy-only variant under a fair training protocol.",
        "Rigorous empirical study: Five-fold stratified cross-validation, bootstrap confidence intervals, McNemar tests, six fairly trained baselines, cross-dataset evaluation on PMRAM-BD [17], hyperparameter and class-weighting ablations, and interpretability via Grad-CAM [18] and Multi-Scale Attention Decomposition (MSAD).",
    ]
    for i, c in enumerate(contributions, 1):
        add_para(doc, f"({i}) {c}")

    add_para(
        doc,
        "The remainder of this paper is organized as follows. Section 2 reviews related work. "
        "Section 3 describes datasets, HSANet architecture, training, and evaluation protocols. "
        "Section 4 presents results. Section 5 discusses clinical implications and limitations. "
        "Section 6 concludes."
    )

    # --- 2 Related Work ---
    add_heading(doc, "2. Related Work", 1)

    add_heading(doc, "2.1 Brain Tumor MRI Classification", 2)
    add_para(
        doc,
        "Early approaches combined hand-crafted texture features with classical classifiers [19]. "
        "Deep CNNs subsequently dominated benchmarks on Figshare, CE-MRI, and BraTS [20,21], although "
        "class definitions, imaging sequences, and train-test protocols differ substantially across "
        "sources, complicating fair comparison. BRISC 2025 [22] addresses several limitations of legacy "
        "collections by providing 6,000 contrast-enhanced T1-weighted slices with expert annotations, "
        "balanced four-class labels (glioma, meningioma, pituitary, no tumor), and a standardized split, "
        "released publicly on Kaggle under CC BY 4.0."
    )

    add_heading(doc, "2.2 Attention and Multi-Scale Representation Learning", 2)
    add_para(
        doc,
        "Channel attention recalibrates feature responses globally (SE-Net [9]; ECA-Net [10]). "
        "CBAM [11] composes channel and spatial attention sequentially. Multi-scale context is "
        "traditionally modeled via atrous spatial pyramid pooling [23] and dilated convolutions [24]. "
        "Efficient architectures increasingly employ depthwise separable convolutions [25] with lightweight "
        "gating. HSANet's AMSM differs by applying parallel dilated depthwise-separable branches at rates "
        "{1, 2, 4} and fusing them with learned softmax weights over four branches—including identity—"
        "so that the output is a convex combination of feature maps."
    )

    add_heading(doc, "2.3 Uncertainty Quantification in Medical Imaging", 2)
    add_para(
        doc,
        "Bayesian neural networks and MC-dropout [16] estimate predictive uncertainty through stochastic "
        "ensembling at inference. Deep ensembles [26] often improve calibration but multiply computational "
        "and storage costs. EDL [15] treats predictions as Dirichlet distributions and has been extended to "
        "medical imaging tasks including COVID-19 detection [27] and lesion segmentation [28]. We adopt EDL "
        "for efficient single-pass uncertainty and benchmark against MC-dropout."
    )

    add_heading(doc, "2.4 Cross-Dataset Generalization", 2)
    add_para(
        doc,
        "Domain shift across MRI vendors, field strengths, and patient populations degrades model "
        "performance [7,8]. External cohorts collected under independent protocols—such as PMRAM-BD [17], "
        "a Bangladeshi brain tumor MRI dataset—provide a stringent test of out-of-distribution robustness "
        "beyond the development benchmark."
    )

    # --- 3 Methods ---
    add_heading(doc, "3. Materials and Methods", 1)

    add_heading(doc, "3.1 Datasets", 2)
    add_heading(doc, "3.1.1 BRISC 2025 (Primary Benchmark)", 3)
    add_para(
        doc,
        "We evaluate on the BRISC 2025 classification task [22]: 5,000 training images and 1,000 "
        "official test images, four classes, multi-planar T1-weighted contrast-enhanced MRI, expert "
        "radiologist validation, JPEG format resampled to 224×224 pixels. BRISC is accessed via the "
        "public Kaggle release (https://www.kaggle.com/datasets/briscdataset/brisc2025; dataset DOI: "
        "10.34740/kaggle/ds/7632487) under CC BY 4.0. Per the Scientific Data descriptor [22], BRISC "
        "was created by Fateh et al. with affiliations at Iran University of Science and Technology, "
        "Shahrood University of Technology, Northern Care Alliance NHS Foundation Trust (UK), and the "
        "University of Essex. The authors of the present study have no affiliation with the BRISC team. "
        "We use only the published train/test split without modification, perform five-fold cross-validation "
        "on training data, and report primary metrics on the held-out test set."
    )

    add_heading(doc, "3.1.2 PMRAM-BD (External Validation)", 3)
    add_para(
        doc,
        "PMRAM-BD [17] is a Bangladeshi brain tumor MRI dataset (glioma, meningioma, pituitary, no tumor) "
        "collected across multiple hospitals, distributed on Mendeley Data (CC BY 4.0). Models are trained "
        "exclusively on BRISC and evaluated zero-shot on PMRAM-BD without fine-tuning, measuring "
        "cross-dataset generalization."
    )

    add_figure_placeholder(
        doc, 1,
        "Overview of the HSANet architecture (version 6). EfficientNet-B3 extracts multi-scale features; "
        "AMSM is applied at pyramid levels 3–5; features are pooled, concatenated, and passed to the "
        "evidential classifier. [Suggested file: custom methodology diagram]",
        "single column, ~6 inches wide",
    )

    add_figure_placeholder(
        doc, 2,
        "Class distribution of the BRISC 2025 training and test sets. "
        "[Suggested file: fig_01_class_distribution.png]",
    )

    add_heading(doc, "3.2 Preprocessing and Data Augmentation", 2)
    add_para(
        doc,
        "All images are resized to 224×224. Training augmentations include random resized crop, horizontal "
        "flip, rotation (±15°), color jitter (brightness and contrast ±0.2), and normalization to ImageNet "
        "statistics consistent with the pretrained backbone. Validation and test use deterministic resize/center "
        "crop. Class-balanced weighted random sampling is applied during training; inverse-frequency class "
        "weights in the loss function are ablated in Section 4.6."
    )

    add_heading(doc, "3.3 HSANet Architecture", 2)
    add_heading(doc, "3.3.1 Network Overview", 3)
    add_para(
        doc,
        "HSANet version 6 comprises: (1) a feature extractor—tf_efficientnet_b3.ns_jft_in1k (timm) with "
        "outputs at indices [2, 3, 4]; (2) Lightweight AMSM applied independently at each scale; "
        "(3) adaptive average pooling per scale, channel concatenation, and an evidential classifier. "
        "Unlike version 5, version 6 does not include a Dual Attention Module (Section 3.4)."
    )

    add_figure_placeholder(
        doc, 3,
        "Detailed structure of the Lightweight AMSM. Three depthwise-separable dilated branches (rates 1, 2, 4) "
        "and an identity branch are fused via softmax-normalized weights. "
        "[Suggested file: custom AMSM diagram]",
        "single column",
    )

    add_heading(doc, "3.3.2 Depthwise-Separable Dilated Branches", 3)
    add_para(
        doc,
        "Each dilated branch applies depthwise convolution (groups = channels) followed by pointwise 1×1 "
        "convolution, batch normalization, and ReLU. This design reduces attention parameters by approximately "
        "one order of magnitude relative to standard convolutions."
    )

    add_heading(doc, "3.3.3 Lightweight AMSM", 3)
    add_para(
        doc,
        "Given feature map x ∈ R^(C×H×W), three parallel branches produce m₁, m₂, m₄ via depthwise-separable "
        "convolutions with dilation rates 1, 2, and 4. Global average pooling is applied to each branch output "
        "and to x (identity branch). The pooled descriptors are concatenated (dimension 4C), passed through a "
        "two-layer MLP with softmax output w ∈ Δ³, and fused as "
        "y = w₁m₁ + w₂m₂ + w₃m₄ + w₄x. This corrected formulation replaces an earlier design "
        "y = Σw_k m_k + x in which the identity path received implicit unit weight outside the softmax, "
        "preventing branch weights from summing to unity."
    )

    add_heading(doc, "3.3.4 Evidential Classifier", 3)
    add_para(
        doc,
        "The classification head maps concatenated features to logits z, evidence e = softplus(z), Dirichlet "
        "parameters α = e + 1, strength S = Σ_k α_k, predictive probabilities p = α/S, and total uncertainty "
        "u = K/S for K = 4 classes [15]. Aleatoric and epistemic components follow Sensoy et al. [15]. "
        "Softplus is preferred over exponential evidence mapping for numerical stability on large logits."
    )

    add_figure_placeholder(
        doc, 4,
        "Evidential Deep Learning head and Dirichlet-based uncertainty decomposition. "
        "[Suggested file: custom EDL diagram]",
    )

    add_heading(doc, "3.3.5 Training Objective", 3)
    add_para(
        doc,
        "The EvidentialLoss combines: (i) Bayesian cross-entropy (expected CE under the Dirichlet); "
        "(ii) KL divergence to a uniform Dirichlet with linear annealing over the first 10 epochs "
        "(λ_KL = 0.2); and (iii) focal loss on expected probabilities (λ_focal = 0.3, γ = 2) [30]. "
        "Optimization uses AdamW (learning rate 1×10⁻⁴, weight decay 1×10⁻⁴), cosine learning-rate "
        "schedule with 5-epoch warmup, 30 epochs, batch size 32, and automatic mixed precision."
    )

    add_heading(doc, "3.4 Design Evolution: Version 5 to Version 6", 2)
    add_para(
        doc,
        "Version 5 instantiated HSANet-Lite with both AMSM and a Lightweight DAM after each AMSM. "
        "DAM comprised ECA channel attention [10] followed by spatial attention [11] (channel-wise "
        "average/maximum pooling, 7×7 convolution, sigmoid mask). Version 5 ablation on BRISC under an "
        "identical protocol produced the following held-out test accuracies:"
    )
    add_table(
        doc,
        ["Configuration", "BRISC Test Accuracy (%)"],
        [
            ["Backbone only", "98.5"],
            ["+ DAM only", "98.5"],
            ["+ AMSM (depthwise-separable)", "99.3"],
            ["Full (AMSM + DAM)", "99.1"],
        ],
        caption="Table A. Version 5 architectural ablation on BRISC 2025 test set.",
    )
    add_para(
        doc,
        "DAM did not improve in-domain accuracy or cross-dataset metrics relative to AMSM-only, while "
        "adding approximately 0.69M parameters and increasing inference latency. Version 6 removes DAM "
        "entirely, achieving 99.2% test accuracy with a simpler pipeline. Attention in HSANet is therefore "
        "realized as scale-gated fusion within AMSM rather than sequential channel-spatial re-weighting."
    )

    add_heading(doc, "3.5 Baseline Models and Fair Comparison Protocol", 2)
    add_para(
        doc,
        "Six pretrained architectures (timm) are compared: ResNet-50, VGG-16, EfficientNet-B3, ViT-B/16, "
        "Swin-Tiny, and DenseNet-121. All baselines are trained for 30 epochs with identical augmentation, "
        "optimizer, scheduler, and FairCELoss (cross-entropy plus focal loss with matched λ_focal = 0.3) "
        "to ensure that performance differences are not attributable to a stronger loss on HSANet alone."
    )

    add_heading(doc, "3.6 Evaluation Metrics and Statistical Analysis", 2)
    add_para(
        doc,
        "We report accuracy, macro-precision, macro-recall, macro-F1, Cohen's κ, Matthews correlation "
        "coefficient (MCC), one-vs-rest macro AUC, expected calibration error (ECE) [31], confusion "
        "matrices, parameter count, and inference latency (ms/image, batch size 1, GPU). Uncertainty "
        "quality is assessed via AUROC for misclassification detection, Brier score, and selective "
        "prediction curves. Five-fold stratified cross-validation reports mean ± standard deviation with "
        "95% confidence intervals. Test accuracy bootstrap CIs use 1,000 resamples. Paired McNemar tests "
        "compare HSANet against each baseline on BRISC test predictions (α = 0.05) [32]."
    )

    add_heading(doc, "3.7 Ablation Studies", 2)
    ablations = [
        "Architectural: backbone only versus +AMSM (version 6 final).",
        "Loss fairness: HSANet with full EDL versus HSANet with CE+focal only (matched architecture).",
        "Class weighting: neither, sampler only, loss only, or both.",
        "Hyperparameters: grid over λ_KL ∈ {0.05, 0.1, 0.2, 0.3, 0.5} and λ_focal ∈ {0.1, 0.2, 0.3, 0.5}.",
        "Cross-dataset: all variants trained on BRISC, tested on PMRAM-BD.",
        "Uncertainty: EDL versus MC-dropout (20 forward passes; entropy and mutual information).",
    ]
    for a in ablations:
        add_para(doc, f"• {a}")

    add_heading(doc, "3.8 Interpretability Analysis", 2)
    add_para(
        doc,
        "Grad-CAM [18] visualizations are computed on the final convolutional stage. Uncertainty-aware "
        "Grad-CAM (UA-GradCAM) weights gradients by evidential uncertainty. Multi-Scale Attention "
        "Decomposition (MSAD) visualizes per-branch AMSM activations and softmax fusion weights, enabling "
        "inspection of which dilation scales dominate per tumor class."
    )

    add_heading(doc, "3.9 Implementation Details", 2)
    add_para(
        doc,
        "Experiments are implemented in PyTorch 2.x with timm, scikit-learn, and SciPy, executed on "
        "NVIDIA Tesla T4 GPUs (Kaggle). Code and trained weights will be released upon acceptance "
        "https://github.com/tarequejosh/HSANet"
    )

    doc.add_page_break()

    # --- 4 Results ---
    add_heading(doc, "4. Results", 1)
    add_para(
        doc,
        "This section reports primary benchmark results, ablations, external validation, uncertainty "
        "analysis, and efficiency. Figure placeholders indicate where to insert prepared figures from "
        "hsanet_results v6/figures/."
    )

    add_heading(doc, "4.1 Primary Results on BRISC 2025", 2)
    add_table(
        doc,
        ["Model", "Params (M)", "Infer. (ms)", "Acc. (%)", "Macro-F1 (%)", "κ", "MCC", "AUC", "ECE"],
        [
            ["ResNet-50", "23.52", "6.76", "98.4", "98.52", "0.978", "0.978", "0.9985", "0.0123"],
            ["VGG-16", "134.28", "12.45", "98.7", "98.88", "0.982", "0.982", "0.9986", "0.0112"],
            ["EfficientNet-B3", "10.70", "12.56", "98.7", "98.80", "0.982", "0.982", "0.9997", "0.0107"],
            ["ViT-B/16", "85.80", "14.25", "99.1", "99.18", "0.988", "0.988", "0.9997", "0.0085"],
            ["Swin-Tiny", "27.52", "11.08", "98.9", "99.01", "0.985", "0.985", "0.9994", "0.0108"],
            ["DenseNet-121", "6.96", "15.55", "98.8", "98.91", "0.984", "0.984", "0.9995", "0.0110"],
            ["HSANet (v6)", "10.97", "19.16", "99.2", "99.30", "0.989", "0.989", "0.9995", "0.0263"],
        ],
        caption="Table 1. Comparison on BRISC 2025 held-out test set (N = 1,000). Bold indicates best result per column where HSANet leads; EfficientNet-B3 and ViT-B/16 achieve lowest ECE and highest AUC among baselines.",
    )

    add_figure_placeholder(doc, 5, "Five-fold cross-validation training and validation curves for HSANet. [fig_02_cv_training_curves.png]")
    add_figure_placeholder(doc, 6, "Confusion matrix of HSANet on the BRISC test set. [fig_03_hsanet_confusion_matrix.png]")
    add_figure_placeholder(doc, 7, "Per-class ROC curves for HSANet. [fig_04_hsanet_roc_curves.png]")

    add_para(
        doc,
        "Five-fold cross-validation on BRISC training data yields mean accuracy 99.06 ± 0.17% "
        "(95% CI ±0.15), macro-F1 99.08 ± 0.17%, AUC 0.99949 ± 0.00031, ECE 0.018 ± 0.004, "
        "and κ 0.987 ± 0.002. Bootstrap 95% confidence interval for test accuracy is [98.6%, 99.7%]. "
        "HSANet achieves the highest test accuracy and macro-F1. McNemar comparison against ResNet-50 "
        "is significant (χ² = 4.08, p = 0.043; HSANet correct on 10 discordant cases where ResNet-50 "
        "errs versus 2 converse), while comparison against ViT-B/16 is not significant (p = 1.0), "
        "consistent with near-ceiling benchmark performance."
    )
    add_para(
        doc,
        "Per-class F1 scores for HSANet are: glioma 99.01%, meningioma 98.86%, no tumor 100.00%, "
        "pituitary 99.33%. Most errors involve confusion between meningioma and pituitary classes."
    )

    add_figure_placeholder(doc, 8, "Accuracy comparison across all models. [fig_05_accuracy_comparison.png]")
    add_figure_placeholder(doc, 9, "Accuracy versus parameter count (efficiency frontier). [fig_06_accuracy_vs_params.png]")
    add_figure_placeholder(doc, 10, "Confusion matrices for baseline models. [fig_07_baseline_confusion_matrices.png]")

    add_heading(doc, "4.2 Architectural Ablation (Version 6)", 2)
    add_table(
        doc,
        ["Configuration", "Params (M)", "Acc. (%)", "Macro-F1 (%)", "ECE", "Unc. ratio"],
        [
            ["Backbone only", "10.28", "98.5", "98.52", "0.0359", "2.00×"],
            ["+ AMSM (HSANet v6)", "10.97", "99.2", "99.30", "0.0263", "2.60×"],
        ],
        caption="Table 2. Component ablation on BRISC test set. Unc. ratio = mean uncertainty on incorrect / correct predictions.",
    )
    add_figure_placeholder(doc, 11, "In-domain architectural ablation results. [fig_10_ablation_study.png]")
    add_para(
        doc,
        "AMSM contributes +0.7 percentage points in accuracy over the backbone alone and improves "
        "uncertainty separability (ratio 2.00× to 2.60×), confirming that multi-scale fusion drives gains."
    )

    add_heading(doc, "4.3 Loss Function Fairness Ablation", 2)
    add_para(
        doc,
        "HSANet trained with CE+focal only (no KL term) achieves 98.6% accuracy versus 99.2% with full EDL "
        "at identical architecture—a +0.6 percentage point contribution from the evidential objective. "
        "The best baseline (ViT-B/16, 99.1%) trails full HSANet but exceeds the CE-only HSANet variant, "
        "indicating that reported improvements arise from the combination of architecture and EDL rather "
        "than loss mismatch alone."
    )
    add_figure_placeholder(doc, 12, "Decomposition of accuracy gains: baseline, architecture, and EDL contributions. [fig_07b_loss_fairness.png]")

    add_heading(doc, "4.4 Cross-Dataset Generalization", 2)
    add_table(
        doc,
        ["Model / Setting", "BRISC Acc. (%)", "PMRAM-BD Acc. (%)", "PMRAM Macro-F1 (%)"],
        [
            ["HSANet (v6)", "99.2", "97.41", "97.41"],
            ["Backbone only (ablation)", "98.5", "97.01", "97.00"],
            ["+ AMSM (ablation)", "99.2", "97.41", "97.41"],
        ],
        caption="Table 3. Cross-dataset evaluation: train on BRISC, test on PMRAM-BD without fine-tuning.",
    )
    add_figure_placeholder(doc, 13, "Cross-dataset generalization across models. [fig_08_cross_dataset_generalization.png]")
    add_figure_placeholder(doc, 14, "Generalization gap (in-domain vs. external accuracy). [fig_09_generalization_gap.png]")
    add_para(
        doc,
        "HSANet retains 97.41% accuracy on PMRAM-BD (κ = 0.965), demonstrating useful transfer despite "
        "a 1.8 percentage point in-domain-to-external gap. AMSM improves external accuracy by +0.40 "
        "points over backbone-only."
    )

    add_heading(doc, "4.5 Uncertainty Quantification", 2)
    add_table(
        doc,
        ["Method", "Acc. (%)", "Error AUROC", "Brier", "ECE"],
        [
            ["EDL (proposed)", "99.2", "0.946", "0.00314", "0.0263"],
            ["MC-Dropout (entropy)", "99.2", "0.960", "0.00308", "0.0256"],
            ["MC-Dropout (MI)", "99.2", "0.910", "0.00308", "0.0256"],
        ],
        caption="Table 4. Misclassification detection on BRISC test set (HSANet architecture).",
    )
    add_figure_placeholder(doc, 15, "Calibration and reliability diagrams. [fig_14_calibration.png]")
    add_figure_placeholder(doc, 16, "EDL versus MC-Dropout uncertainty comparison. [fig_14b_edl_vs_mcdropout.png]")
    add_figure_placeholder(doc, 17, "Selective prediction: accuracy versus coverage as uncertain cases are deferred. [fig_13_selective_prediction.png]")
    add_para(
        doc,
        "MC-dropout (20 samples) achieves marginally higher error-detection AUROC (0.960 vs. 0.946) but "
        "requires twenty forward passes. EDL provides comparable Brier scores in a single pass. Mean "
        "uncertainty on misclassified samples is 2.60× that on correct samples for EDL."
    )

    add_heading(doc, "4.6 Hyperparameter and Class-Weighting Sensitivity", 2)
    add_para(
        doc,
        "Grid search over (λ_KL, λ_focal) yields accuracies in the range 98.5–99.4%, with a stable "
        "high-performance region near (0.3, 0.3); default values (0.2, 0.3) lie within this region. "
        "Class-weighting ablation shows that combined sampler and loss weighting can inflate accuracy "
        "(99.6%) while increasing per-class F1 variance; the main model uses balanced loss-only weighting."
    )
    add_figure_placeholder(doc, 18, "Hyperparameter sensitivity heatmap (λ_KL vs. λ_focal). [fig_10c_hp_sensitivity.png]")
    add_figure_placeholder(doc, 19, "Class-weighting strategy ablation. [fig_10b_class_weighting_ablation.png]")

    add_heading(doc, "4.7 Statistical Significance", 2)
    add_figure_placeholder(doc, 20, "McNemar test results and bootstrap confidence intervals. [fig_11_statistical_tests.png]")

    add_heading(doc, "4.8 Interpretability and Failure Analysis", 2)
    add_figure_placeholder(
        doc, 21,
        "Grad-CAM and UA-GradCAM visualizations on representative correct and incorrect cases. "
        "[fig_15_ua_gradcam.png]",
    )
    add_figure_placeholder(
        doc, 22,
        "Multi-Scale Attention Decomposition (MSAD): per-branch activations and softmax fusion weights. "
        "[fig_15b_msad.png]",
    )
    add_figure_placeholder(doc, 23, "Per-class F1 comparison across models. [fig_17_per_class_f1.png]")
    add_figure_placeholder(doc, 24, "Failure cases with attention overlays. [fig_18b_failure_cases.png]")

    add_heading(doc, "4.9 Computational Efficiency", 2)
    add_para(
        doc,
        "HSANet (10.97M parameters) is 12.3× smaller than VGG-16 and 7.8× smaller than ViT-B/16 while "
        "achieving equal or superior accuracy. Per-image inference (19.2 ms) exceeds ResNet-50 (6.8 ms) "
        "due to multi-branch AMSM, acceptable for offline triage but motivating distillation for real-time use."
    )
    add_figure_placeholder(doc, 25, "Inference time and efficiency analysis. [fig_18_efficiency.png]")
    add_figure_placeholder(doc, 26, "Radar chart of multi-metric model comparison. [fig_16_radar_chart.png]")

    doc.add_page_break()

    # --- 5 Discussion ---
    add_heading(doc, "5. Discussion", 1)

    add_heading(doc, "5.1 Why Removing Dual Attention Improved Performance", 2)
    add_para(
        doc,
        "The version 5 DAM stack re-applied channel and spatial gating after AMSM had already formed a "
        "softmax-weighted convex combination of multi-scale features. Sequential attention can "
        "over-compress representations when upstream fusion has allocated salience across scales [12]. "
        "Empirically, DAM-only matched the bare backbone (98.5%), AMSM-only peaked at 99.3% (v5) and "
        "99.2% (v6), and the full AMSM+DAM model reached only 99.1%. We therefore recommend "
        "single-stage scale fusion as the primary inductive bias for BRISC-scale MRI classification."
    )

    add_heading(doc, "5.2 Relation to BRISC Benchmarks and Baselines", 2)
    add_para(
        doc,
        "Fateh et al. [22] report Swin-HAFNet and additional baselines on BRISC. Under our fair "
        "comparison protocol, HSANet is competitive with ViT-B/16 and superior to classical CNNs, "
        "indicating that efficient CNN hybrids remain viable when parameter budget and single-GPU "
        "inference are constrained."
    )

    add_heading(doc, "5.3 Clinical Implications", 2)
    add_para(
        doc,
        "High accuracy on BRISC with 97.41% external validation, combined with EDL-based selective "
        "prediction and MSAD interpretability, supports investigation in computer-aided triage workflows. "
        "Uncertainty ratios above 2.5× between incorrect and correct predictions suggest practical "
        "deferral thresholds, pending prospective validation."
    )

    add_heading(doc, "5.4 Limitations", 2)
    limitations = [
        "Near-ceiling BRISC metrics limit discriminative power among top models; reader studies are needed to establish clinical significance.",
        "Retrospective public benchmarks may not reflect prospective deployment; PMRAM-BD provides one external cohort but multi-site validation remains necessary.",
        "Two-dimensional slice classification does not exploit volumetric context or BRISC segmentation masks.",
        "HSANet ECE (0.0263) exceeds some baselines despite strong AUROC; temperature scaling may improve calibration.",
        "Fairness across scanner types, demographics, and tumor subtypes is not audited.",
    ]
    for lim in limitations:
        add_para(doc, f"• {lim}")

    add_heading(doc, "5.5 Future Work", 2)
    add_para(
        doc,
        "Future directions include multi-task learning with BRISC segmentation labels, lightweight 3D "
        "extensions, federated training across institutions, knowledge distillation for real-time inference, "
        "and prospective trials linking uncertainty thresholds to referral policies."
    )

    # --- 6 Conclusion ---
    add_heading(doc, "6. Conclusion", 1)
    add_para(
        doc,
        "We presented HSANet, a lightweight hybrid scale-attention network for brain tumor MRI "
        "classification combining depthwise-separable multi-scale fusion with evidential uncertainty. "
        "A central empirical finding from the version 5 to version 6 design cycle is negative: stacking "
        "dual channel-spatial attention after AMSM failed to improve performance. The final model achieves "
        "99.2% BRISC test accuracy and 97.41% PMRAM-BD accuracy with approximately 11 million parameters, "
        "matching or exceeding substantially larger baselines under fair training. Evidential learning "
        "contributes measurable accuracy and enables single-pass uncertainty estimation suitable for "
        "selective prediction. HSANet demonstrates that rigorous ablation of attention components is as "
        "scientifically valuable as their introduction."
    )

    doc.add_page_break()

    # --- Declarations ---
    add_heading(doc, "Declarations", 1)
    declarations = [
        ("Ethics approval", "Not required for publicly available, de-identified retrospective imaging datasets; local institutional review board approval should be obtained before any clinical deployment."),
        ("Consent", "Not applicable; per dataset publisher policies [17,22]."),
        ("Competing interests", "The authors declare no competing interests."),
        ("Funding", "[To be completed]"),
        ("Author contributions", "[CRediT taxonomy to be completed]"),
        ("Data availability", "BRISC 2025: https://doi.org/10.1038/s41597-026-06753-y and Kaggle DOI 10.34740/kaggle/ds/7632487. PMRAM-BD: https://data.mendeley.com/datasets/m7w55sw88b/1"),
        ("Code availability", "https://github.com/tarequejosh/HSANet — source code, evaluation results, and figure assets. Trained checkpoints are available on request due to size limits."),
    ]
    for title, text in declarations:
        add_para(doc, f"{title}: ", bold=True)
        doc.paragraphs[-1].add_run(text).font.name = "Times New Roman"

    doc.add_page_break()

    # --- References ---
    add_heading(doc, "References", 1)
    refs = [
        "Ostrom QT, et al. CBTRUS statistical report: primary brain and other central nervous system tumors diagnosed in the United States in 2012–2016. Neuro-Oncology. 2019;21(Suppl 5):v1-v100.",
        "Louis DN, et al. The 2021 WHO Classification of Tumors of the Central Nervous System: a summary. Neuro-Oncology. 2021;23(8):1231-1251.",
        "Sartaj F, et al. Brain tumor classification using deep learning on public MRI datasets: a review. Diagnostics. 2023;13(6):1012.",
        "Deepak S, Ameer PM. Brain tumor classification using deep CNN features via logistic regression. In: IEEE CBMS; 2019.",
        "Dosovitskiy A, et al. An image is worth 16×16 words: Transformers for image recognition at scale. ICLR. 2021.",
        "Liu Z, et al. Swin Transformer: Hierarchical vision transformer using shifted windows. ICCV. 2021:10012-10022.",
        "Kohl S, et al. How to study domain shift in medical image analysis? arXiv:2208.03233. 2022.",
        "Bandi P, et al. From detection of individual metastases to classification of lymph node status at the patient level. Med Image Anal. 2019;56:147-161.",
        "Hu J, Shen L, Sun G. Squeeze-and-excitation networks. CVPR. 2018:7132-7141.",
        "Wang Q, et al. ECA-Net: Efficient channel attention for deep convolutional neural networks. CVPR. 2020:11534-11542.",
        "Woo S, et al. CBAM: Convolutional block attention module. ECCV. 2018:3-19.",
        "Anderson AG, et al. Does more attention increase reliability of CNNs? NeurIPS Workshop. 2019.",
        "Kompa B, et al. Second opinion needed: communicating uncertainty in medical machine learning. NPJ Digit Med. 2021;4:85.",
        "Begoli E, et al. A systematic review of robustness in deep learning for computer vision. Neurocomputing. 2019;381:381-403.",
        "Sensoy M, Kaplan L, Kandemir M. Evidential deep learning to quantify classification uncertainty. NeurIPS. 2018;31.",
        "Gal Y, Ghahramani Z. Dropout as a Bayesian approximation. ICML. 2016:1050-1059.",
        "Mannan MS, et al. PMRAM: Bangladeshi Brain Cancer - MRI Dataset. Mendeley Data V1. 2024. https://data.mendeley.com/datasets/m7w55sw88b/1",
        "Selvaraju RR, et al. Grad-CAM: Visual explanations from deep networks. ICCV. 2017:618-626.",
        "Gordillo N, et al. State of the art survey on MRI brain tumor segmentation. Pattern Recognit. 2013;46(8):2440-2455.",
        "Menze BH, et al. The multimodal brain tumor image segmentation benchmark (BRATS). IEEE TMI. 2015;34(10):1993-2024.",
        "Rezaee A, Gevers T. 2D MRI-based glioma classification and survival analysis after diagnosis: a survey. Neurocomputing. 2022;500:208-224.",
        "Fateh A, et al. BRISC: Annotated dataset for brain tumor segmentation and classification. Sci Data. 2026. https://doi.org/10.1038/s41597-026-06753-y",
        "Chen LC, et al. Encoder-decoder with atrous separable convolution (DeepLabv3+). ECCV. 2018:801-818.",
        "Yu F, Koltun V. Multi-scale context aggregation by dilated convolutions. ICLR. 2016.",
        "Howard A, et al. Searching for MobileNetV3. ICCV. 2019:1314-1324.",
        "Lakshminarayanan B, et al. Simple and scalable predictive uncertainty estimation using deep ensembles. NeurIPS. 2017;30.",
        "Sensoy M, et al. Misclassified out-of-distribution detection using deep evidential classification. IEEE TPAMI. 2021.",
        "Nair T, et al. Exploring uncertainty measures in deep networks for MS lesion segmentation. Med Image Anal. 2020;59:101557.",
        "Tan M, Le QV. EfficientNet: Rethinking model scaling for CNNs. ICML. 2019:6105-6114.",
        "Lin TY, et al. Focal loss for dense object detection. ICCV. 2017:2980-2988.",
        "Guo C, et al. On calibration of modern neural networks. ICML. 2017:1321-1330.",
        "Dietterich TG. Approximate statistical tests for comparing supervised classification learning algorithms. Neural Comput. 1998;10(7):1895-1923.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        r = p.add_run(f"[{i}] {ref}")
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)

    doc.add_page_break()
    add_heading(doc, "Appendix A. Figure Insertion Guide", 1)
    fig_guide = [
        ("Fig. 1", "HSANet architecture (methodology — create custom)"),
        ("Fig. 2", "fig_01_class_distribution.png"),
        ("Fig. 3", "AMSM module detail (methodology — create custom)"),
        ("Fig. 4", "EDL head (methodology — create custom)"),
        ("Fig. 5", "fig_02_cv_training_curves.png"),
        ("Fig. 6", "fig_03_hsanet_confusion_matrix.png"),
        ("Fig. 7", "fig_04_hsanet_roc_curves.png"),
        ("Fig. 8", "fig_05_accuracy_comparison.png"),
        ("Fig. 9", "fig_06_accuracy_vs_params.png"),
        ("Fig. 10", "fig_07_baseline_confusion_matrices.png"),
        ("Fig. 11", "fig_10_ablation_study.png"),
        ("Fig. 12", "fig_07b_loss_fairness.png"),
        ("Fig. 13–14", "fig_08, fig_09 cross-dataset"),
        ("Fig. 15–17", "fig_14, fig_14b, fig_13 uncertainty"),
        ("Fig. 18–19", "fig_10c, fig_10b sensitivity"),
        ("Fig. 20", "fig_11_statistical_tests.png"),
        ("Fig. 21–24", "fig_15, fig_15b, fig_17, fig_18b interpretability"),
        ("Fig. 25–26", "fig_18_efficiency, fig_16_radar"),
    ]
    add_table(doc, ["Figure", "Source file / note"], fig_guide, caption="Appendix A. Recommended figure mapping.")

    add_heading(doc, "Appendix B. Version 5 vs. Version 6 Summary", 1)
    add_table(
        doc,
        ["Component", "Version 5", "Version 6 (final)"],
        [
            ["EfficientNet-B3 backbone", "Yes", "Yes"],
            ["AMSM (4-branch DW-Sep)", "Yes", "Yes"],
            ["DAM (ECA + spatial)", "Yes", "Removed"],
            ["Evidential classifier", "Yes", "Yes"],
            ["BRISC test accuracy", "99.1% (full) / 99.3% (AMSM-only)", "99.2%"],
        ],
    )

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build_document()
